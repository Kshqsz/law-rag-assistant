# coding: utf-8
"""
API 路由模块：文档上传与管理接口
"""
import os
import uuid
import shutil
from typing import List
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File
from sqlalchemy.orm import Session

from ..database import get_db, Document, User
from ..schemas import DocumentResponse, DocumentListResponse, SuccessResponse
from ..auth import get_current_user

router = APIRouter(prefix="/documents", tags=["文档管理"])

# 文档上传目录
UPLOAD_DIR = "./uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# 支持的文件类型
ALLOWED_EXTENSIONS = {".txt", ".md", ".pdf", ".doc", ".docx"}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB


def get_file_extension(filename: str) -> str:
    """获取文件扩展名"""
    return os.path.splitext(filename)[1].lower()


def validate_file(file: UploadFile) -> None:
    """验证上传文件"""
    ext = get_file_extension(file.filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"不支持的文件类型: {ext}，支持的类型: {', '.join(ALLOWED_EXTENSIONS)}"
        )


@router.post("/upload", response_model=DocumentResponse, summary="上传文档")
async def upload_document(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    上传文档
    
    支持的文件类型: .txt, .md, .pdf, .doc, .docx
    最大文件大小: 10MB
    """
    # 验证文件类型
    validate_file(file)
    
    # 生成唯一文件名
    ext = get_file_extension(file.filename)
    unique_filename = f"{uuid.uuid4().hex}{ext}"
    
    # 用户专属目录
    user_dir = os.path.join(UPLOAD_DIR, str(current_user.id))
    os.makedirs(user_dir, exist_ok=True)
    
    file_path = os.path.join(user_dir, unique_filename)
    
    # 保存文件
    try:
        content = await file.read()
        
        # 检查文件大小
        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"文件大小超过限制 ({MAX_FILE_SIZE // 1024 // 1024}MB)"
            )
        
        with open(file_path, "wb") as f:
            f.write(content)
        
        file_size = len(content)
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"文件保存失败: {str(e)}"
        )
    
    # 保存到数据库
    document = Document(
        user_id=current_user.id,
        filename=unique_filename,
        original_filename=file.filename,
        file_path=file_path,
        file_size=file_size,
        file_type=ext,
        is_processed=False
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    
    return document


@router.get("", response_model=DocumentListResponse, summary="获取文档列表")
async def list_documents(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """获取当前用户上传的所有文档"""
    documents = db.query(Document).filter(
        Document.user_id == current_user.id
    ).order_by(Document.created_at.desc()).all()
    
    return {"documents": documents}


@router.get("/{document_id}", response_model=DocumentResponse, summary="获取文档详情")
async def get_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """获取指定文档的详情"""
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="文档不存在"
        )
    
    return document


@router.get("/{document_id}/content", summary="获取文档内容")
async def get_document_content(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """获取文档的文本内容"""
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="文档不存在"
        )
    
    try:
        if document.file_type in [".txt", ".md"]:
            with open(document.file_path, "r", encoding="utf-8") as f:
                content = f.read()
        elif document.file_type == ".pdf":
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(document.file_path)
                text_parts = []
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                content = '\n'.join(text_parts) if text_parts else "无法提取 PDF 内容"
            except ImportError:
                content = "服务端未安装 PyPDF2，无法解析 PDF 文件"
        elif document.file_type in [".docx"]:
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(document.file_path)
                text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
                content = '\n'.join(text_parts) if text_parts else "无法提取 Word 内容"
            except ImportError:
                content = "服务端未安装 python-docx，无法解析 Word 文件"
        elif document.file_type == ".doc":
            content = "暂不支持 .doc 格式，请转换为 .docx 后上传"
        else:
            content = f"暂不支持预览 {document.file_type} 格式的文件内容"
        
        return {"content": content}
        
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"读取文件失败: {str(e)}"
        )


@router.delete("/{document_id}", response_model=SuccessResponse, summary="删除文档")
async def delete_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """删除指定文档"""
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="文档不存在"
        )
    
    # 删除文件
    try:
        if os.path.exists(document.file_path):
            os.remove(document.file_path)
    except Exception as e:
        pass  # 忽略文件删除错误
    
    # 从数据库删除
    db.delete(document)
    db.commit()
    
    return {"success": True, "message": "文档已删除"}


@router.post("/{document_id}/process", response_model=DocumentResponse, summary="处理文档入库")
async def process_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    将文档处理并导入向量库
    
    注意：此功能会将文档内容向量化并存入 ChromaDB，
    以便在问答时能够检索到相关内容
    """
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="文档不存在"
        )
    
    if document.is_processed:
        return document
    
    try:
        # 1. 读取文档内容
        content = None
        if document.file_type in [".txt", ".md"]:
            with open(document.file_path, "r", encoding="utf-8") as f:
                content = f.read()
        elif document.file_type == ".pdf":
            from PyPDF2 import PdfReader
            reader = PdfReader(document.file_path)
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            content = '\n'.join(text_parts)
        elif document.file_type in [".docx"]:
            from docx import Document as DocxDocument
            doc = DocxDocument(document.file_path)
            text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
            content = '\n'.join(text_parts)
        
        if not content:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="无法提取文档内容"
            )
        
        # 2. 分割文档
        from langchain.text_splitter import RecursiveCharacterTextSplitter
        from langchain.schema import Document as LangDocument
        
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            separators=["\n\n", "\n", "。", "；", " ", ""]
        )
        
        lang_doc = LangDocument(
            page_content=content,
            metadata={"source": document.original_filename, "doc_id": str(document.id)}
        )
        chunks = splitter.split_documents([lang_doc])
        
        # 3. 向量化并存入 ChromaDB
        from law_ai.utils import get_embedding_model, get_vectorstore
        
        embedder = get_embedding_model()
        vectorstore = get_vectorstore("user_docs")
        
        texts = [chunk.page_content for chunk in chunks]
        metadatas = [chunk.metadata for chunk in chunks]
        vectorstore.add_texts(texts=texts, metadatas=metadatas, embedding=embedder)
        
        # 4. 标记为已处理
        document.is_processed = True
        db.commit()
        db.refresh(document)
        
        return document
        
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"文档处理失败: {str(e)}"
        )
